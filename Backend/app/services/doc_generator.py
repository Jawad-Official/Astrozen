import logging
import io
from app.services.ai_service import AiService
from sqlalchemy.orm import Session
from app.models.project_idea import ProjectIdea
from app.models.project_asset import ProjectAsset, AssetType, AssetStatus
from app.services.storage_service import R2StorageProvider
from uuid import UUID
from typing import List, Dict, Any, Optional
from fpdf import FPDF
from docx import Document
from docx.shared import Pt
import markdown2
import re

logger = logging.getLogger(__name__)

class DocGenerator:
    def __init__(self, ai_service: AiService, storage: R2StorageProvider):
        self.ai_service = ai_service
        self.storage = storage

    async def get_doc_questions(self, db: Session, idea_id: UUID, doc_type: AssetType) -> List[str]:
        idea = db.query(ProjectIdea).filter(ProjectIdea.id == idea_id).first()
        if not idea:
            return []

        # Get previous docs as context
        prev_docs = db.query(ProjectAsset).filter(
            ProjectAsset.project_idea_id == idea_id,
            ProjectAsset.generation_status == AssetStatus.COMPLETE
        ).all()
        
        context = f"Project: {idea.refined_description}\n"
        for doc in prev_docs:
            context += f"Previous Doc ({doc.asset_type}): [Content included in prompt]\n"

        system_prompt = f"You are a technical architect. Based on the project idea and previous docs, ask 2-3 essential questions to generate a high-quality {doc_type}."
        user_prompt = f"Context:\n{context}\n\nWhat extra info do you need for {doc_type}?"
        
        try:
            # In a real app, this would be a JSON prompt. For now, we expect a bulleted list.
            response = await self.ai_service.generate_text(system_prompt, user_prompt)
            questions = [q.strip("- ").strip() for q in response.split("\n") if q.strip()]
            return questions[:3] # Limit to 3
        except Exception as e:
            logger.error(f"Failed to get questions for {doc_type}: {e}")
            return ["What are the specific technical constraints?", "Who are the primary end-users?"]

    async def generate_single_doc(self, db: Session, idea_id: UUID, doc_type: AssetType, user_answers: Optional[Dict[str, str]] = None):
        idea = db.query(ProjectIdea).filter(ProjectIdea.id == idea_id).first()
        if not idea:
            return

        asset = db.query(ProjectAsset).filter(
            ProjectAsset.project_idea_id == idea_id,
            ProjectAsset.asset_type == doc_type
        ).first()
        
        if not asset:
            asset = ProjectAsset(project_idea_id=idea_id, asset_type=doc_type)
            db.add(asset)
        
        asset.generation_status = AssetStatus.GENERATING
        db.commit()

        # Get context from all sources
        context = f"Original Idea: {idea.refined_description}\n\n"
        
        if idea.validation_report:
            report = idea.validation_report
            context += "--- Project Validation Report ---\n"
            context += f"Market Analysis: {report.market_analysis}\n"
            context += f"Core Features: {report.core_features}\n"
            context += f"Tech Stack: {report.tech_stack}\n"
            context += f"Pricing Model: {report.pricing_model}\n\n"

        # Context from all previous documents
        all_docs = db.query(ProjectAsset).filter(
            ProjectAsset.project_idea_id == idea_id,
            ProjectAsset.generation_status == AssetStatus.COMPLETE
        ).all()
        for d in all_docs:
            context += f"--- {d.asset_type} ---\n[Context from {d.asset_type} included]\n"
        
        system_prompt = f"You are a technical writer. Generate a comprehensive {doc_type} in Markdown."
        if not user_answers:
            system_prompt += " The user has skipped questions, so use your best technical judgment for defaults."
        
        user_prompt = f"Context:\n{context}\n\nUser Input: {user_answers}\n\nGenerate {doc_type} Content:"
        
        try:
            content_md = await self.ai_service.generate_text(system_prompt, user_prompt)
            
            # Paths
            base_path = f"projects/{idea_id}/docs/{doc_type.value}"
            md_path = f"{base_path}.md"
            pdf_path = f"{base_path}.pdf"
            docx_path = f"{base_path}.docx"

            # 1. Upload Markdown
            self.storage.upload_file(content_md.encode('utf-8'), md_path, content_type="text/markdown")

            # 2. Convert and Upload PDF
            pdf_bytes = self._generate_pdf(content_md, f"{doc_type.value.replace('_', ' ')}")
            self.storage.upload_file(pdf_bytes, pdf_path, content_type="application/pdf")

            # 3. Convert and Upload DOCX
            docx_bytes = self._generate_docx(content_md, f"{doc_type.value.replace('_', ' ')}")
            self.storage.upload_file(docx_bytes, docx_path, content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

            # Update asset record
            asset.storage_path = md_path
            asset.storage_path_pdf = pdf_path
            asset.storage_path_docx = docx_path
            asset.file_format = "md"
            asset.generation_status = AssetStatus.COMPLETE
            
        except Exception as e:
            logger.error(f"Failed to generate {doc_type}: {e}")
            asset.generation_status = AssetStatus.FAILED
        
        db.commit()
        return asset

    def _generate_pdf(self, md_content: str, title: str) -> bytes:
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("helvetica", "B", 16)
        pdf.cell(0, 10, title, ln=True, align="C")
        pdf.ln(5)
        
        pdf.set_font("helvetica", size=11)
        # Clean markdown characters for simple PDF rendering
        clean_text = re.sub(r'#+\s', '', md_content)
        clean_text = re.sub(r'\*\*', '', clean_text)
        
        # Split by lines and add to PDF
        for line in clean_text.split('\n'):
            pdf.multi_cell(0, 8, line)
        
        return pdf.output()

    def _generate_docx(self, md_content: str, title: str) -> bytes:
        doc = Document()
        doc.add_heading(title, 0)
        
        # Simple markdown to docx conversion
        for line in md_content.split('\n'):
            if line.startswith('### '):
                doc.add_heading(line.replace('### ', ''), level=2)
            elif line.startswith('## '):
                doc.add_heading(line.replace('## ', ''), level=1)
            elif line.startswith('# '):
                doc.add_heading(line.replace('# ', ''), level=0)
            elif line.strip():
                p = doc.add_paragraph(line)
                p.style.font.size = Pt(11)
        
        target_stream = io.BytesIO()
        doc.save(target_stream)
        return target_stream.getvalue()

    async def generate_all_docs(self, db: Session, idea_id: UUID):
        """Used for background processing if the user wants all at once (fallback)"""
        for doc_type in AssetType:
            if doc_type != AssetType.DIAGRAM_MERMAID:
                await self.generate_single_doc(db, idea_id, doc_type)